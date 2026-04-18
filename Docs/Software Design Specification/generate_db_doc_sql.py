#!/usr/bin/env python3
"""
Generate a Word (.docx) document describing every database table
from the SQL DDL schema file (schema.sql).

Usage:
    cd "D:\\FPTU\\Capstone\\Code\\BE\\Core\\Docs\\Software Design Specification"
    py generate_db_doc_sql.py

Output: database-design-tables.docx
"""

import re, os, sys
from collections import OrderedDict

# ---------------------------------------------------------------------------
# Auto-install python-docx if missing
# ---------------------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SCHEMA_PATH = os.path.join(SCRIPT_DIR, "schema.sql")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "database-design-tables.docx")

# ---------------------------------------------------------------------------
# Domain groups - ordered
# ---------------------------------------------------------------------------
DOMAIN_GROUPS = OrderedDict([
    ("Core Tables", [
        "customer", "dish_category", "dish", "restaurant_table",
        "reservation", "reservation_table",
    ]),
    ("Order Tables", [
        "orders", "order_item", "payment", "tax",
        "order_coupon", "order_promotion",
    ]),
    ("Inventory Tables", [
        "supplier", "ingredient", "current_stock",
        "ingredient_supplier", "recipe",
        "inventory_transaction", "inventory_transaction_item",
    ]),
    ("Staff & Security Tables", [
        "staff_account", "role", "permission", "role_permission",
        "auth_session", "login_activity", "audit_log",
        "service_error", "service_error_category",
    ]),
    ("Promotion & Coupon Tables", [
        "promotion", "promotion_rule", "promotion_target",
        "coupon",
    ]),
    ("Shift & Attendance Tables", [
        "shift_template", "shift_assignment",
        "attendance_record", "time_log",
    ]),
    ("Notification Tables", [
        "notifications", "notification_read_states",
        "notification_preferences",
    ]),
    ("Infrastructure Tables", [
        "lookup_type", "lookup_value",
        "media_asset", "dish_media", "dish_tag",
        "table_media", "inventory_transaction_media",
        "i18n_language", "i18n_text", "i18n_translation",
        "system_setting", "email_template",
    ]),
])

# ---------------------------------------------------------------------------
# Table descriptions
# ---------------------------------------------------------------------------
DESCRIPTIONS = {
    "customer": "Stores customer profiles including contact info, membership status, and loyalty points.",
    "dish_category": "Groups dishes into categories for menu organization and display ordering.",
    "dish": "Stores individual dish/menu item details including pricing, descriptions, and preparation info.",
    "restaurant_table": "Defines physical dining tables with capacity, zone, type, status, and QR code info.",
    "reservation": "Records customer table reservations with party size, time, status, and source channel.",
    "reservation_table": "Junction table linking reservations to specific restaurant tables (M:N).",
    "orders": "Central order entity tracking table, customer, staff, amounts, tax, and status.",
    "order_item": "Individual line items within an order, linking a dish with quantity and price.",
    "payment": "Records payment transactions for orders including method, amount received, and change.",
    "tax": "Defines tax rate configurations (inclusive/exclusive) applied to orders.",
    "order_coupon": "Junction table applying a coupon discount to a specific order.",
    "order_promotion": "Junction table applying a promotion discount to a specific order.",
    "supplier": "Stores supplier/vendor contact and business information for ingredient sourcing.",
    "ingredient": "Defines raw ingredients used in recipes, with category, type, unit, and i18n support.",
    "current_stock": "Tracks real-time on-hand stock quantity and minimum threshold per ingredient.",
    "ingredient_supplier": "Junction table linking ingredients to their suppliers (M:N).",
    "recipe": "Defines the ingredient composition of each dish with quantity and unit.",
    "inventory_transaction": "Header record for inventory operations (import, export, stock-check) with approval workflow.",
    "inventory_transaction_item": "Line items within an inventory transaction specifying ingredient, quantity, and variance.",
    "staff_account": "Stores staff user accounts with credentials, role assignment, and account status.",
    "role": "Defines staff roles (e.g., Admin, Manager, Chef, Waiter) with a configurable status.",
    "permission": "Defines granular permissions as screen-action pairs for RBAC.",
    "role_permission": "Junction table assigning permissions to roles (M:N).",
    "auth_session": "Tracks active authentication sessions (JWT tokens) for staff accounts.",
    "login_activity": "Logs login/logout events for security auditing and session tracking.",
    "audit_log": "Records staff actions on data entities for compliance and change tracking.",
    "service_error": "Logs service quality incidents linked to staff, orders, items, or tables.",
    "service_error_category": "Categorizes service error types for classification and reporting.",
    "promotion": "Defines time-bound promotional offers with discount value, usage limits, and status.",
    "promotion_rule": "Defines eligibility conditions for a promotion (min order value, required dish/category).",
    "promotion_target": "Specifies which dishes or categories a promotion applies to.",
    "coupon": "Defines customer-specific or general coupon codes with discount and validity period.",
    "shift_template": "Defines reusable shift time templates with start/end times and buffer periods.",
    "shift_assignment": "Assigns a staff member to a specific shift on a given work date.",
    "attendance_record": "Records actual check-in/out times, lateness, early leave, and review status per shift.",
    "time_log": "Granular punch-in/punch-out log entries with GPS and device validation per attendance.",
    "notifications": "Stores system-generated notification messages with targeting and priority metadata.",
    "notification_read_states": "Tracks per-user read and acknowledgment status for each notification.",
    "notification_preferences": "Stores per-user notification type preferences (enabled, sound settings).",
    "lookup_type": "Defines categories of lookup enumerations (e.g., OrderStatus, PaymentMethod, Zone).",
    "lookup_value": "Individual enum/lookup values belonging to a lookup type, with soft-delete support.",
    "media_asset": "Stores uploaded media file metadata (URL, MIME type, dimensions, duration).",
    "dish_media": "Junction table linking dishes to their media assets (images/videos).",
    "dish_tag": "Junction table assigning lookup-based tags to dishes for filtering.",
    "table_media": "Junction table linking restaurant tables to their media assets.",
    "inventory_transaction_media": "Junction table attaching media evidence to inventory transactions.",
    "i18n_language": "Defines supported languages for the internationalization system.",
    "i18n_text": "Stores source text entries keyed for translation, with context hints.",
    "i18n_translation": "Stores translated text for each (text_id, lang_code) pair.",
    "system_setting": "Key-value configuration store with typed values and sensitivity flags.",
    "email_template": "Stores HTML email templates for system-generated emails.",
}

# ---------------------------------------------------------------------------
# Field descriptions - keyed by (table_name, column_name)
# ---------------------------------------------------------------------------
FIELD_DESCRIPTIONS = {
    # -- customer --
    ("customer", "customer_id"): "Auto-incremented unique identifier for the customer",
    ("customer", "full_name"): "Customer's full name",
    ("customer", "phone"): "Customer's phone number (unique contact identifier)",
    ("customer", "email"): "Customer's email address",
    ("customer", "is_member"): "Whether the customer has registered as a member",
    ("customer", "loyalty_points"): "Accumulated loyalty reward points",
    ("customer", "created_at"): "Timestamp when the customer record was created",

    # -- dish_category --
    ("dish_category", "category_id"): "Auto-incremented unique identifier for the category",
    ("dish_category", "category_name"): "Display name of the dish category",
    ("dish_category", "category_name_text_id"): "FK to i18n_text for category name translation",
    ("dish_category", "description"): "Short description of the category",
    ("dish_category", "description_text_id"): "FK to i18n_text for description translation",
    ("dish_category", "display_order"): "Sort position for menu display",
    ("dish_category", "is_disable"): "Whether the category is disabled/hidden from the menu",

    # -- dish --
    ("dish", "dish_id"): "Auto-incremented unique identifier for the dish",
    ("dish", "category_id"): "FK to dish_category this dish belongs to",
    ("dish", "dish_name"): "Display name of the dish",
    ("dish", "price"): "Current selling price of the dish",
    ("dish", "created_at"): "Timestamp when the dish was created",
    ("dish", "dish_status_lv_id"): "FK to lookup_value for dish status (Active, Inactive, etc.)",
    ("dish", "description"): "Full description of the dish",
    ("dish", "slogan"): "Marketing slogan for the dish",
    ("dish", "note"): "Internal notes about the dish",
    ("dish", "calories"): "Calorie count per serving",
    ("dish", "short_description"): "Brief description for list/card views",
    ("dish", "display_order"): "Sort position within the category",
    ("dish", "chef_recommended"): "Whether the dish is chef-recommended",
    ("dish", "prep_time_minutes"): "Preparation time in minutes",
    ("dish", "cook_time_minutes"): "Cooking time in minutes",
    ("dish", "isOnline"): "Whether the dish is available for online ordering",
    ("dish", "description_text_id"): "FK to i18n_text for description translation",
    ("dish", "slogan_text_id"): "FK to i18n_text for slogan translation",
    ("dish", "note_text_id"): "FK to i18n_text for note translation",
    ("dish", "short_description_text_id"): "FK to i18n_text for short description translation",
    ("dish", "dish_name_text_id"): "FK to i18n_text for dish name translation",

    # -- restaurant_table --
    ("restaurant_table", "table_id"): "Auto-incremented unique identifier for the table",
    ("restaurant_table", "table_code"): "Unique display code for the table (e.g., T01, VIP-A)",
    ("restaurant_table", "capacity"): "Maximum seating capacity of the table",
    ("restaurant_table", "table_qr_img"): "FK to media_asset for the table's QR code image",
    ("restaurant_table", "table_status_lv_id"): "FK to lookup_value for table status (Available, Occupied, etc.)",
    ("restaurant_table", "table_type_lv_id"): "FK to lookup_value for table type (Regular, VIP, etc.)",
    ("restaurant_table", "zone_lv_id"): "FK to lookup_value for table zone/area (Indoor, Outdoor, etc.)",
    ("restaurant_table", "isOnline"): "Whether the table is visible for online reservation",
    ("restaurant_table", "qr_token"): "Unique token embedded in the table's QR code for ordering",
    ("restaurant_table", "created_at"): "Timestamp when the table record was created",
    ("restaurant_table", "updated_at"): "Timestamp of the last update",
    ("restaurant_table", "updated_by_staff_id"): "FK to staff_account who last updated this table",
    ("restaurant_table", "is_deleted"): "Soft-delete flag; 1 = table is logically deleted",

    # -- reservation --
    ("reservation", "reservation_id"): "Auto-incremented unique identifier for the reservation",
    ("reservation", "customer_id"): "FK to customer; NULL if guest reservation",
    ("reservation", "customer_name"): "Name provided for the reservation",
    ("reservation", "phone"): "Contact phone for the reservation",
    ("reservation", "email"): "Contact email for the reservation",
    ("reservation", "party_size"): "Number of guests in the party",
    ("reservation", "reserved_time"): "Date and time of the reservation",
    ("reservation", "created_at"): "Timestamp when the reservation was created",
    ("reservation", "source_lv_id"): "FK to lookup_value for booking source (Walk-in, Phone, Online)",
    ("reservation", "notes"): "Additional notes or special requests",
    ("reservation", "reservation_status_lv_id"): "FK to lookup_value for reservation status",

    # -- reservation_table --
    ("reservation_table", "reservation_id"): "FK to reservation (composite PK part 1)",
    ("reservation_table", "table_id"): "FK to restaurant_table (composite PK part 2)",

    # -- orders --
    ("orders", "order_id"): "Auto-incremented unique identifier for the order",
    ("orders", "table_id"): "FK to restaurant_table; NULL for takeout/delivery orders",
    ("orders", "staff_id"): "FK to staff_account who created the order",
    ("orders", "created_at"): "Timestamp when the order was created",
    ("orders", "customer_id"): "FK to customer placing the order",
    ("orders", "updated_at"): "Timestamp of the last status or amount update",
    ("orders", "total_amount"): "Final total amount after tax and discounts",
    ("orders", "tip_amount"): "Optional tip amount added to the order",
    ("orders", "source_lv_id"): "FK to lookup_value for order source (Dine-in, Online, POS)",
    ("orders", "order_status_lv_id"): "FK to lookup_value for order status (Pending, Confirmed, etc.)",
    ("orders", "tax_id"): "FK to tax configuration applied to this order",
    ("orders", "tax_amount"): "Calculated tax amount",
    ("orders", "sub_total_amount"): "Subtotal before tax and discounts",

    # -- order_item --
    ("order_item", "order_item_id"): "Auto-incremented unique identifier for the order item",
    ("order_item", "order_id"): "FK to orders this item belongs to",
    ("order_item", "dish_id"): "FK to dish being ordered",
    ("order_item", "quantity"): "Number of units ordered",
    ("order_item", "price"): "Unit price at time of order",
    ("order_item", "item_status"): "Legacy status code (1=Created, 2=InProgress, 3=Ready, 4=Served, 5=Rejected)",
    ("order_item", "reject_reason"): "Reason for rejection if item was rejected",
    ("order_item", "Note"): "Special instructions or notes for this item",
    ("order_item", "item_status_lv_id"): "FK to lookup_value for item status (replaces legacy item_status)",

    # -- payment --
    ("payment", "payment_id"): "Auto-incremented unique identifier for the payment",
    ("payment", "order_id"): "FK to orders being paid for",
    ("payment", "received_amount"): "Amount received from the customer",
    ("payment", "change_amount"): "Change returned to the customer",
    ("payment", "paid_at"): "Timestamp when the payment was made",
    ("payment", "method_lv_id"): "FK to lookup_value for payment method (Cash, Card, etc.)",

    # -- tax --
    ("tax", "tax_id"): "Auto-incremented unique identifier for the tax configuration",
    ("tax", "tax_name"): "Display name of the tax (e.g., VAT 8%)",
    ("tax", "tax_rate"): "Tax rate as a percentage (e.g., 8.00 for 8%)",
    ("tax", "tax_type"): "Whether tax is INCLUSIVE or EXCLUSIVE of listed prices",
    ("tax", "is_active"): "Whether this tax configuration is currently active",
    ("tax", "is_default"): "Whether this is the default tax applied to new orders",
    ("tax", "created_at"): "Timestamp when the tax was created",
    ("tax", "updated_at"): "Timestamp of the last update",

    # -- order_coupon --
    ("order_coupon", "order_coupon_id"): "Auto-incremented unique identifier",
    ("order_coupon", "order_id"): "FK to orders receiving the coupon discount",
    ("order_coupon", "coupon_id"): "FK to coupon being applied",
    ("order_coupon", "discount_amount"): "Actual discount amount applied to the order",
    ("order_coupon", "applied_at"): "Timestamp when the coupon was applied",

    # -- order_promotion --
    ("order_promotion", "order_promotion_id"): "Auto-incremented unique identifier",
    ("order_promotion", "order_id"): "FK to orders receiving the promotion discount",
    ("order_promotion", "promotion_id"): "FK to promotion being applied",
    ("order_promotion", "discount_amount"): "Actual discount amount applied to the order",
    ("order_promotion", "applied_at"): "Timestamp when the promotion was applied",

    # -- supplier --
    ("supplier", "supplier_id"): "Auto-incremented unique identifier for the supplier",
    ("supplier", "supplier_name"): "Business name of the supplier",
    ("supplier", "phone"): "Supplier's contact phone number",
    ("supplier", "email"): "Supplier's contact email address",
    ("supplier", "address"): "Supplier's business address",
    ("supplier", "tax_code"): "Supplier's tax identification number",

    # -- ingredient --
    ("ingredient", "ingredient_id"): "Auto-incremented unique identifier for the ingredient",
    ("ingredient", "ingredient_name"): "Display name of the ingredient",
    ("ingredient", "type_lv_id"): "FK to lookup_value for ingredient type classification",
    ("ingredient", "ingredient_name_text_id"): "FK to i18n_text for ingredient name translation",
    ("ingredient", "image_id"): "FK to media_asset for ingredient image",
    ("ingredient", "unit_lv_id"): "FK to lookup_value for default measurement unit",
    ("ingredient", "category_lv_id"): "FK to lookup_value for ingredient category",

    # -- current_stock --
    ("current_stock", "ingredient_id"): "FK to ingredient (also serves as PK - 1:1 relationship)",
    ("current_stock", "quantity_on_hand"): "Current available stock quantity",
    ("current_stock", "last_updated_at"): "Timestamp of the last stock quantity change",
    ("current_stock", "min_stock_level"): "Minimum stock threshold for low-stock alerts",

    # -- ingredient_supplier --
    ("ingredient_supplier", "ingredient_supplier_id"): "Auto-incremented unique identifier",
    ("ingredient_supplier", "supplier_id"): "FK to supplier providing the ingredient",
    ("ingredient_supplier", "ingredient_id"): "FK to ingredient supplied",
    ("ingredient_supplier", "created_at"): "Timestamp when the link was created",

    # -- recipe --
    ("recipe", "dish_id"): "FK to dish (composite PK part 1)",
    ("recipe", "ingredient_id"): "FK to ingredient (composite PK part 2)",
    ("recipe", "quantity"): "Amount of ingredient required per dish serving",
    ("recipe", "unit"): "Measurement unit for the quantity",
    ("recipe", "note"): "Preparation notes for this ingredient in the recipe",

    # -- inventory_transaction --
    ("inventory_transaction", "transaction_id"): "Auto-incremented unique identifier for the transaction",
    ("inventory_transaction", "created_by"): "FK to staff_account who created the transaction",
    ("inventory_transaction", "created_at"): "Timestamp when the transaction was created",
    ("inventory_transaction", "note"): "General notes or description for the transaction",
    ("inventory_transaction", "type_lv_id"): "FK to lookup_value for transaction type (Import, Export, StockCheck)",
    ("inventory_transaction", "status_lv_id"): "FK to lookup_value for transaction status (Draft, Submitted, Approved)",
    ("inventory_transaction", "supplier_id"): "FK to supplier for import transactions",
    ("inventory_transaction", "approved_at"): "Timestamp when the transaction was approved",
    ("inventory_transaction", "approved_by"): "FK to staff_account who approved the transaction",
    ("inventory_transaction", "export_reason_lv_id"): "FK to lookup_value for export reason (if type is Export)",
    ("inventory_transaction", "stock_check_area_note"): "Area/zone note for stock-check transactions",
    ("inventory_transaction", "submitted_at"): "Timestamp when the transaction was submitted for approval",
    ("inventory_transaction", "transaction_code"): "Human-readable transaction code (e.g., IMP-20240101-001)",

    # -- inventory_transaction_item --
    ("inventory_transaction_item", "transaction_item_id"): "Auto-incremented unique identifier for the line item",
    ("inventory_transaction_item", "transaction_id"): "FK to inventory_transaction header",
    ("inventory_transaction_item", "ingredient_id"): "FK to ingredient being transacted",
    ("inventory_transaction_item", "quantity"): "Planned/requested quantity",
    ("inventory_transaction_item", "note"): "Notes specific to this line item",
    ("inventory_transaction_item", "actual_quantity"): "Actual counted quantity (for stock-check)",
    ("inventory_transaction_item", "system_quantity"): "System-recorded quantity at time of stock-check",
    ("inventory_transaction_item", "unit_lv_id"): "FK to lookup_value for measurement unit",
    ("inventory_transaction_item", "unit_price"): "Unit price for import transactions",
    ("inventory_transaction_item", "variance_reason_lv_id"): "FK to lookup_value for stock variance reason",

    # -- staff_account --
    ("staff_account", "account_id"): "Auto-incremented unique identifier for the staff account",
    ("staff_account", "full_name"): "Staff member's full name",
    ("staff_account", "phone"): "Staff member's phone number",
    ("staff_account", "email"): "Staff member's email address",
    ("staff_account", "role_id"): "FK to role assigned to this staff member",
    ("staff_account", "created_at"): "Timestamp when the account was created",
    ("staff_account", "username"): "Unique login username",
    ("staff_account", "password_hash"): "Bcrypt-hashed password (never stored in plain text)",
    ("staff_account", "is_locked"): "Whether the account is locked (login disabled)",
    ("staff_account", "last_login_at"): "Timestamp of the most recent successful login",
    ("staff_account", "RegisteredDeviceId"): "Device ID registered for push notifications",
    ("staff_account", "account_status_lv_id"): "FK to lookup_value for account status (Active, Suspended, etc.)",

    # -- role --
    ("role", "role_id"): "Auto-incremented unique identifier for the role",
    ("role", "role_code"): "Unique code identifying the role (e.g., ADMIN, MANAGER)",
    ("role", "role_name"): "Display name of the role",
    ("role", "role_status_lv_id"): "FK to lookup_value for role status",

    # -- permission --
    ("permission", "permission_id"): "Auto-incremented unique identifier for the permission",
    ("permission", "screen_code"): "Screen/page code the permission applies to",
    ("permission", "action_code"): "Action code (e.g., VIEW, CREATE, EDIT, DELETE)",

    # -- role_permission --
    ("role_permission", "role_id"): "FK to role (composite PK part 1)",
    ("role_permission", "permission_id"): "FK to permission (composite PK part 2)",

    # -- auth_session --
    ("auth_session", "session_id"): "Auto-incremented unique identifier for the session",
    ("auth_session", "user_id"): "FK to staff_account owning this session",
    ("auth_session", "token_hash"): "SHA-256 hash of the refresh token",
    ("auth_session", "expires_at"): "Expiration timestamp of the session/token",
    ("auth_session", "created_at"): "Timestamp when the session was created",
    ("auth_session", "revoked"): "Whether the session has been explicitly revoked",
    ("auth_session", "device_info"): "User-Agent or device information string",
    ("auth_session", "ip_address"): "IP address from which the session was created",

    # -- login_activity --
    ("login_activity", "login_activity_id"): "Auto-incremented unique identifier",
    ("login_activity", "staff_id"): "FK to staff_account involved in the login event",
    ("login_activity", "session_id"): "FK to auth_session; NULL if session was not created",
    ("login_activity", "event_type"): "Type of login event (Login, Logout, FailedLogin, TokenRefresh)",
    ("login_activity", "device_info"): "User-Agent or device info at time of event",
    ("login_activity", "ip_address"): "IP address at time of event",
    ("login_activity", "occurred_at"): "Timestamp when the event occurred",

    # -- audit_log --
    ("audit_log", "log_id"): "Auto-incremented unique identifier for the audit entry",
    ("audit_log", "staff_id"): "FK to staff_account who performed the action",
    ("audit_log", "action_code"): "Code describing the action (e.g., CREATE, UPDATE, DELETE)",
    ("audit_log", "target_table"): "Name of the database table affected",
    ("audit_log", "target_id"): "Primary key of the affected record",
    ("audit_log", "created_at"): "Timestamp when the action was logged",

    # -- service_error --
    ("service_error", "error_id"): "Auto-incremented unique identifier for the service error",
    ("service_error", "staff_id"): "FK to staff_account who caused or reported the error",
    ("service_error", "order_id"): "FK to orders related to the error (optional)",
    ("service_error", "order_item_id"): "FK to order_item related to the error (optional)",
    ("service_error", "table_id"): "FK to restaurant_table related to the error (optional)",
    ("service_error", "category_id"): "FK to service_error_category for error classification",
    ("service_error", "description"): "Detailed description of the service error incident",
    ("service_error", "penalty_amount"): "Monetary penalty assessed for the error",
    ("service_error", "is_resolved"): "Whether the error has been resolved",
    ("service_error", "resolved_by"): "FK to staff_account who resolved the error",
    ("service_error", "resolved_at"): "Timestamp when the error was resolved",
    ("service_error", "created_at"): "Timestamp when the error was reported",
    ("service_error", "severity_lv_id"): "FK to lookup_value for error severity (Low, Medium, High, Critical)",

    # -- service_error_category --
    ("service_error_category", "category_id"): "Auto-incremented unique identifier",
    ("service_error_category", "category_code"): "Unique code for the error category",
    ("service_error_category", "category_name"): "Display name of the error category",
    ("service_error_category", "description"): "Description of what this category covers",
    ("service_error_category", "category_name_text_id"): "FK to i18n_text for category name translation",
    ("service_error_category", "category_desc_text_id"): "FK to i18n_text for description translation",

    # -- promotion --
    ("promotion", "promotion_id"): "Auto-incremented unique identifier for the promotion",
    ("promotion", "promo_code"): "Optional promotional code string",
    ("promotion", "promo_name"): "Display name of the promotion",
    ("promotion", "description"): "Description of the promotion offer",
    ("promotion", "start_time"): "Start date/time when the promotion becomes active",
    ("promotion", "end_time"): "End date/time when the promotion expires",
    ("promotion", "discount_value"): "Discount amount or percentage value",
    ("promotion", "max_usage"): "Maximum number of times this promotion can be used",
    ("promotion", "used_count"): "Number of times this promotion has been applied",
    ("promotion", "created_at"): "Timestamp when the promotion was created",
    ("promotion", "type_lv_id"): "FK to lookup_value for promotion type (Percentage, FixedAmount)",
    ("promotion", "promotion_status_lv_id"): "FK to lookup_value for promotion status (Active, Expired, etc.)",
    ("promotion", "promo_name_text_id"): "FK to i18n_text for promotion name translation",
    ("promotion", "promo_desc_text_id"): "FK to i18n_text for promotion description translation",

    # -- promotion_rule --
    ("promotion_rule", "rule_id"): "Auto-incremented unique identifier for the rule",
    ("promotion_rule", "promotion_id"): "FK to promotion this rule belongs to",
    ("promotion_rule", "min_order_value"): "Minimum order value to qualify for the promotion",
    ("promotion_rule", "min_quantity"): "Minimum item quantity to qualify",
    ("promotion_rule", "required_dish_id"): "FK to dish required in order to qualify",
    ("promotion_rule", "required_category_id"): "FK to dish_category required in order to qualify",

    # -- promotion_target --
    ("promotion_target", "target_id"): "Auto-incremented unique identifier",
    ("promotion_target", "promotion_id"): "FK to promotion this target belongs to",
    ("promotion_target", "dish_id"): "FK to specific dish the promotion applies to",
    ("promotion_target", "category_id"): "FK to dish_category the promotion applies to",

    # -- coupon --
    ("coupon", "coupon_id"): "Auto-incremented unique identifier for the coupon",
    ("coupon", "coupon_code"): "Unique coupon code string for redemption",
    ("coupon", "coupon_name"): "Display name of the coupon",
    ("coupon", "description"): "Description of the coupon offer",
    ("coupon", "start_time"): "Start date/time when the coupon becomes valid",
    ("coupon", "end_time"): "End date/time when the coupon expires",
    ("coupon", "discount_value"): "Discount amount or percentage value",
    ("coupon", "max_usage"): "Maximum number of times this coupon can be redeemed",
    ("coupon", "used_count"): "Number of times this coupon has been redeemed",
    ("coupon", "created_at"): "Timestamp when the coupon was created",
    ("coupon", "type_lv_id"): "FK to lookup_value for coupon type (Percentage, FixedAmount)",
    ("coupon", "coupon_status_lv_id"): "FK to lookup_value for coupon status",
    ("coupon", "customer_id"): "FK to customer if coupon is customer-specific; NULL for general",

    # -- shift_template --
    ("shift_template", "shift_template_id"): "Auto-incremented unique identifier for the shift template",
    ("shift_template", "template_name"): "Display name of the shift template (e.g., Morning, Evening)",
    ("shift_template", "default_start_time"): "Default start time for shifts using this template",
    ("shift_template", "default_end_time"): "Default end time for shifts using this template",
    ("shift_template", "description"): "Description of the shift template",
    ("shift_template", "buffer_before_minutes"): "Minutes of buffer allowed before shift start for early check-in",
    ("shift_template", "buffer_after_minutes"): "Minutes of buffer allowed after shift end for late check-out",
    ("shift_template", "is_active"): "Whether the template is currently active and available for assignment",
    ("shift_template", "created_by"): "FK to staff_account who created the template",
    ("shift_template", "created_at"): "Timestamp when the template was created",
    ("shift_template", "updated_by"): "FK to staff_account who last updated the template",
    ("shift_template", "updated_at"): "Timestamp of the last update",

    # -- shift_assignment --
    ("shift_assignment", "shift_assignment_id"): "Auto-incremented unique identifier for the assignment",
    ("shift_assignment", "shift_template_id"): "FK to shift_template defining the time pattern",
    ("shift_assignment", "staff_id"): "FK to staff_account assigned to this shift",
    ("shift_assignment", "work_date"): "Calendar date of the assigned shift",
    ("shift_assignment", "planned_start_at"): "Planned start date/time (may differ from template default)",
    ("shift_assignment", "planned_end_at"): "Planned end date/time (may differ from template default)",
    ("shift_assignment", "assignment_status_lv_id"): "FK to lookup_value for assignment status (Scheduled, Published, etc.)",
    ("shift_assignment", "is_active"): "Whether the assignment is currently active",
    ("shift_assignment", "tags"): "Comma-separated tags for categorization",
    ("shift_assignment", "notes"): "Additional notes for the assignment",
    ("shift_assignment", "assigned_by"): "FK to staff_account who created the assignment",
    ("shift_assignment", "assigned_at"): "Timestamp when the assignment was made",
    ("shift_assignment", "created_at"): "Timestamp when the record was created",
    ("shift_assignment", "updated_at"): "Timestamp of the last update",

    # -- attendance_record --
    ("attendance_record", "attendance_id"): "Auto-incremented unique identifier for the attendance record",
    ("attendance_record", "shift_assignment_id"): "FK to shift_assignment this attendance belongs to",
    ("attendance_record", "attendance_status_lv_id"): "FK to lookup_value for attendance status (Present, Absent, Late, etc.)",
    ("attendance_record", "actual_check_in_at"): "Actual check-in timestamp",
    ("attendance_record", "actual_check_out_at"): "Actual check-out timestamp",
    ("attendance_record", "late_minutes"): "Number of minutes the staff was late",
    ("attendance_record", "early_leave_minutes"): "Number of minutes the staff left early",
    ("attendance_record", "worked_minutes"): "Total minutes actually worked",
    ("attendance_record", "is_manual_adjustment"): "Whether the record was manually adjusted by a manager",
    ("attendance_record", "adjustment_reason"): "Reason for manual adjustment",
    ("attendance_record", "reviewed_by"): "FK to staff_account who reviewed/approved the attendance",
    ("attendance_record", "reviewed_at"): "Timestamp when the attendance was reviewed",
    ("attendance_record", "created_at"): "Timestamp when the attendance record was created",
    ("attendance_record", "updated_at"): "Timestamp of the last update",

    # -- time_log --
    ("time_log", "time_log_id"): "Auto-incremented unique identifier for the time log entry",
    ("time_log", "attendance_record_id"): "FK to attendance_record this punch belongs to",
    ("time_log", "punch_in_time"): "Timestamp of the punch-in action",
    ("time_log", "punch_out_time"): "Timestamp of the punch-out action",
    ("time_log", "gps_location_in"): "GPS coordinates at punch-in (lat,lng format)",
    ("time_log", "gps_location_out"): "GPS coordinates at punch-out (lat,lng format)",
    ("time_log", "device_id_in"): "Device identifier used for punch-in",
    ("time_log", "device_id_out"): "Device identifier used for punch-out",
    ("time_log", "validation_status"): "Validation result of the punch (Valid, InvalidLocation, etc.)",
    ("time_log", "punch_duration_minutes"): "Duration in minutes between punch-in and punch-out",
    ("time_log", "created_at"): "Timestamp when the time log entry was created",

    # -- notifications --
    ("notifications", "notification_id"): "Auto-incremented unique identifier for the notification",
    ("notifications", "type"): "Notification type code (e.g., ORDER_CREATED, SHIFT_ASSIGNED)",
    ("notifications", "title"): "Notification title displayed to the user",
    ("notifications", "body"): "Notification body/message content",
    ("notifications", "priority"): "Priority level (Low, Normal, High, Critical)",
    ("notifications", "require_ack"): "Whether the notification requires explicit acknowledgment",
    ("notifications", "sound_key"): "Sound file key to play when notification is received",
    ("notifications", "action_url"): "URL to navigate to when the notification is clicked",
    ("notifications", "entity_type"): "Type of entity the notification relates to (Order, Shift, etc.)",
    ("notifications", "entity_id"): "ID of the related entity",
    ("notifications", "metadata_json"): "Additional metadata in JSON format",
    ("notifications", "target_permissions"): "Comma-separated permission codes to target specific roles",
    ("notifications", "target_user_ids"): "Comma-separated user IDs to target specific users",
    ("notifications", "created_at"): "Timestamp when the notification was created",

    # -- notification_read_states --
    ("notification_read_states", "notification_read_state_id"): "Auto-incremented unique identifier",
    ("notification_read_states", "notification_id"): "FK to notifications this read state belongs to",
    ("notification_read_states", "user_id"): "FK to staff_account (the recipient user)",
    ("notification_read_states", "is_read"): "Whether the user has read the notification",
    ("notification_read_states", "is_acknowledged"): "Whether the user has acknowledged the notification",
    ("notification_read_states", "read_at"): "Timestamp when the notification was read",
    ("notification_read_states", "acknowledged_at"): "Timestamp when the notification was acknowledged",
    ("notification_read_states", "created_at"): "Timestamp when the read state record was created",

    # -- notification_preferences --
    ("notification_preferences", "notification_preference_id"): "Auto-incremented unique identifier",
    ("notification_preferences", "user_id"): "FK to staff_account this preference belongs to",
    ("notification_preferences", "notification_type"): "Notification type code this preference applies to",
    ("notification_preferences", "is_enabled"): "Whether notifications of this type are enabled for the user",
    ("notification_preferences", "sound_enabled"): "Whether sound is enabled for this notification type",
    ("notification_preferences", "created_at"): "Timestamp when the preference was created",
    ("notification_preferences", "updated_at"): "Timestamp of the last preference update",

    # -- lookup_type --
    ("lookup_type", "type_id"): "Auto-incremented unique identifier for the lookup type",
    ("lookup_type", "type_code"): "Unique code for the lookup type (e.g., ORDER_STATUS, PAYMENT_METHOD)",
    ("lookup_type", "type_name"): "Display name of the lookup type",
    ("lookup_type", "description"): "Description of what this lookup type represents",
    ("lookup_type", "is_configurable"): "1 = admin can add/remove values; 0 = controlled enum",
    ("lookup_type", "is_system"): "1 = system-defined type; 0 = user-defined/custom type",
    ("lookup_type", "type_name_text_id"): "FK to i18n_text for type name translation",
    ("lookup_type", "type_desc_text_id"): "FK to i18n_text for description translation",

    # -- lookup_value --
    ("lookup_value", "value_id"): "Auto-incremented unique identifier for the lookup value",
    ("lookup_value", "type_id"): "FK to lookup_type this value belongs to",
    ("lookup_value", "value_code"): "Unique code within the type (e.g., PENDING, CONFIRMED)",
    ("lookup_value", "value_name"): "Display name of the lookup value",
    ("lookup_value", "sort_order"): "Sort position within the type for display ordering",
    ("lookup_value", "is_active"): "Whether the value is currently active and selectable",
    ("lookup_value", "meta"): "Optional JSON metadata for additional configuration",
    ("lookup_value", "is_system"): "1 = system/seeded value; 0 = user-added value",
    ("lookup_value", "locked"): "1 = value_code cannot be changed and value cannot be deleted",
    ("lookup_value", "deleted_at"): "Soft-delete timestamp; never hard-delete lookup values",
    ("lookup_value", "description"): "Description of this specific lookup value",
    ("lookup_value", "update_at"): "Timestamp of the last update",
    ("lookup_value", "value_name_text_id"): "FK to i18n_text for value name translation",
    ("lookup_value", "value_desc_text_id"): "FK to i18n_text for description translation",

    # -- media_asset --
    ("media_asset", "media_id"): "Auto-incremented unique identifier for the media asset",
    ("media_asset", "url"): "Public URL or storage path of the media file",
    ("media_asset", "mime_type"): "MIME type of the file (e.g., image/jpeg, video/mp4)",
    ("media_asset", "width"): "Width in pixels (for images/videos)",
    ("media_asset", "height"): "Height in pixels (for images/videos)",
    ("media_asset", "duration_sec"): "Duration in seconds (for audio/video files)",
    ("media_asset", "created_at"): "Timestamp when the media was uploaded",
    ("media_asset", "media_type_lv_id"): "FK to lookup_value for media type classification",

    # -- dish_media --
    ("dish_media", "dish_id"): "FK to dish (composite PK part 1)",
    ("dish_media", "media_id"): "FK to media_asset (composite PK part 2)",
    ("dish_media", "is_primary"): "Whether this is the primary/thumbnail media for the dish",

    # -- dish_tag --
    ("dish_tag", "dish_tag_id"): "Auto-incremented unique identifier",
    ("dish_tag", "dish_id"): "FK to dish being tagged",
    ("dish_tag", "tag_id"): "FK to lookup_value representing the tag",

    # -- table_media --
    ("table_media", "table_id"): "FK to restaurant_table (composite PK part 1)",
    ("table_media", "media_id"): "FK to media_asset (composite PK part 2)",
    ("table_media", "is_primary"): "Whether this is the primary media for the table",

    # -- inventory_transaction_media --
    ("inventory_transaction_media", "transaction_id"): "FK to inventory_transaction (composite PK part 1)",
    ("inventory_transaction_media", "media_id"): "FK to media_asset (composite PK part 2)",
    ("inventory_transaction_media", "is_primary"): "Whether this is the primary evidence media",

    # -- i18n_language --
    ("i18n_language", "lang_code"): "ISO language code (e.g., en, vi, zh) - primary key",
    ("i18n_language", "lang_name"): "Display name of the language (e.g., English, Tieng Viet)",
    ("i18n_language", "is_active"): "Whether the language is currently active in the system",

    # -- i18n_text --
    ("i18n_text", "text_id"): "Auto-incremented unique identifier for the text entry",
    ("i18n_text", "text_key"): "Unique key identifying the text (e.g., dish.name.pho_bo)",
    ("i18n_text", "source_lang_code"): "FK to i18n_language; language of the source text",
    ("i18n_text", "source_text"): "Original source text content",
    ("i18n_text", "context"): "Context hint to help translators understand usage",
    ("i18n_text", "created_at"): "Timestamp when the text entry was created",
    ("i18n_text", "updated_at"): "Timestamp of the last update",

    # -- i18n_translation --
    ("i18n_translation", "text_id"): "FK to i18n_text (composite PK part 1)",
    ("i18n_translation", "lang_code"): "FK to i18n_language (composite PK part 2)",
    ("i18n_translation", "translated_text"): "Translated text content in the target language",
    ("i18n_translation", "updated_at"): "Timestamp of the last translation update",

    # -- system_setting --
    ("system_setting", "setting_id"): "Auto-incremented unique identifier for the setting",
    ("system_setting", "setting_key"): "Unique configuration key (e.g., RESTAURANT_NAME, MAX_TABLES)",
    ("system_setting", "setting_name"): "Human-readable name of the setting",
    ("system_setting", "value_type"): "Data type of the setting value (STRING, INT, DECIMAL, BOOL, JSON, DATETIME)",
    ("system_setting", "value_string"): "String-typed value storage",
    ("system_setting", "value_int"): "Integer-typed value storage",
    ("system_setting", "value_decimal"): "Decimal-typed value storage",
    ("system_setting", "value_bool"): "Boolean-typed value storage",
    ("system_setting", "value_json"): "JSON-typed value storage",
    ("system_setting", "description"): "Description of what this setting controls",
    ("system_setting", "is_sensitive"): "Whether the value is sensitive (masked in UI)",
    ("system_setting", "updated_at"): "Timestamp of the last value change",
    ("system_setting", "updated_by"): "FK to staff_account who last changed the setting",

    # -- email_template --
    ("email_template", "TemplateId"): "Auto-incremented unique identifier for the email template",
    ("email_template", "TemplateCode"): "Unique code identifying the template (e.g., WELCOME, ORDER_CONFIRM)",
    ("email_template", "TemplateName"): "Display name of the email template",
    ("email_template", "Subject"): "Email subject line (may contain placeholder tokens)",
    ("email_template", "BodyHtml"): "HTML body of the email template (may contain placeholder tokens)",
    ("email_template", "Description"): "Description of when/how this template is used",
    ("email_template", "CreatedAt"): "Timestamp when the template was created",
    ("email_template", "UpdatedAt"): "Timestamp of the last template update",
}

# Unique columns (business logic)
UNIQUE_COLUMNS = {
    "customer": {"phone"},
    "staff_account": {"username"},
    "role": {"role_code"},
    "lookup_type": {"type_code"},
    "coupon": {"coupon_code"},
    "restaurant_table": {"table_code"},
    "email_template": {"TemplateCode"},
    "system_setting": {"setting_key"},
    "i18n_text": {"text_key"},
}


# ===========================================================================
# SQL DDL Parser
# ===========================================================================
def parse_sql(sql_path):
    """
    Parse a MySQL DDL script and return:
      tables: OrderedDict of table_name -> list of column dicts
      pk_map: dict of table_name -> set of pk column names
      fk_map: dict of table_name -> dict of col_name -> ref_table
    """
    with open(sql_path, "r", encoding="utf-8") as f:
        sql = f.read()

    tables = OrderedDict()
    pk_map = {}
    fk_map = {}

    # -- Parse CREATE TABLE blocks --
    create_re = re.compile(
        r"CREATE\s+TABLE\s+`?(\w+)`?\s*\((.*?)\)\s*ENGINE",
        re.DOTALL | re.IGNORECASE,
    )

    for m in create_re.finditer(sql):
        table_name = m.group(1)
        body = m.group(2)

        columns = []
        pk_cols = set()

        # Extract PRIMARY KEY
        pk_m = re.search(r"PRIMARY\s+KEY\s*\(([^)]+)\)", body, re.IGNORECASE)
        if pk_m:
            for pk_col in pk_m.group(1).split(","):
                pk_cols.add(pk_col.strip().strip("`"))

        # Parse each column line
        for line in body.split("\n"):
            line = line.strip().rstrip(",")
            if not line:
                continue
            # Skip PRIMARY KEY line, KEY/INDEX lines, CONSTRAINT lines
            if re.match(r"(PRIMARY\s+KEY|KEY\s|INDEX\s|UNIQUE\s|CONSTRAINT\s)", line, re.IGNORECASE):
                continue

            # Check for enum type first
            enum_m = re.match(
                r"`?(\w+)`?\s+(enum\s*\([^)]+\))(.*)",
                line,
                re.IGNORECASE,
            )
            if enum_m:
                col_name = enum_m.group(1)
                raw_type = enum_m.group(2).strip()
                rest = enum_m.group(3).strip()
            else:
                # Standard column: col_name type [UNSIGNED] [CHARACTER SET...] ...
                col_m = re.match(
                    r"`?(\w+)`?\s+"
                    r"(\w+(?:\([^)]*\))?(?:\s+UNSIGNED)?)"
                    r"(.*)",
                    line,
                    re.IGNORECASE,
                )
                if not col_m:
                    continue
                col_name = col_m.group(1)
                raw_type = col_m.group(2).strip()
                rest = col_m.group(3).strip()

            # Determine NOT NULL
            not_null = bool(re.search(r"\bNOT\s+NULL\b", rest, re.IGNORECASE))
            auto_inc = bool(re.search(r"\bAUTO_INCREMENT\b", rest, re.IGNORECASE))
            has_default = bool(re.search(r"\bDEFAULT\b", rest, re.IGNORECASE))

            # Extract COMMENT
            comment = ""
            comment_m = re.search(r"COMMENT\s+'([^']*)'", rest, re.IGNORECASE)
            if comment_m:
                comment = comment_m.group(1)

            # Clean up type display
            col_type = clean_type(raw_type)

            columns.append({
                "name": col_name,
                "type": col_type,
                "not_null": not_null,
                "auto_increment": auto_inc,
                "has_default": has_default,
                "comment": comment,
                "is_pk": col_name in pk_cols,
            })

        tables[table_name] = columns
        pk_map[table_name] = pk_cols
        fk_map[table_name] = {}

    # -- Parse ALTER TABLE ... FOREIGN KEY --
    fk_re = re.compile(
        r"ALTER\s+TABLE\s+`?(\w+)`?\s+"
        r"ADD\s+CONSTRAINT\s+\w+\s+"
        r"FOREIGN\s+KEY\s*\(`?(\w+)`?\)\s*"
        r"REFERENCES\s+`?(\w+)`?",
        re.IGNORECASE,
    )
    for m in fk_re.finditer(sql):
        table_name = m.group(1)
        col_name = m.group(2)
        ref_table = m.group(3)
        if table_name in fk_map:
            fk_map[table_name][col_name] = ref_table

    return tables, pk_map, fk_map


def clean_type(raw):
    """Clean up a MySQL type for display."""
    t = raw.strip()
    # Remove CHARACTER SET ... COLLATE ...
    t = re.sub(r"\s+CHARACTER\s+SET\s+\w+\s+COLLATE\s+\w+", "", t, flags=re.IGNORECASE)
    # Normalize spacing
    t = re.sub(r"\s+", " ", t).strip()
    return t


# ===========================================================================
# Word Document Generation
# ===========================================================================
HEADER_COLOR = "FBE4D5"  # Peach
FONT_NAME = "Times New Roman"
TABLE_SECTION_PREFIX = "1.3"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top="single", bottom="single", left="single", right="single", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="{top}" w:sz="{sz}" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="{bottom}" w:sz="{sz}" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="{left}" w:sz="{sz}" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="{right}" w:sz="{sz}" w:space="0" w:color="000000"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(borders)


def format_cell(cell, text, bold=False, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    # Reduce paragraph spacing
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)


def generate_docx(tables, pk_map, fk_map):
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)

    # Title
    title_para = doc.add_heading("Database Design - Table Descriptions", level=1)
    for run in title_para.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(16)

    table_sequence = 0
    for domain_name, table_list in DOMAIN_GROUPS.items():
        # Domain heading
        h2 = doc.add_heading(domain_name, level=2)
        for run in h2.runs:
            run.font.name = FONT_NAME

        for tbl_name in table_list:
            if tbl_name not in tables:
                print(f"  WARNING: Table '{tbl_name}' not found in SQL schema, skipping.")
                continue
            table_sequence += 1
            cols = tables[tbl_name]
            fks = fk_map.get(tbl_name, {})
            pks = pk_map.get(tbl_name, set())
            uniques = UNIQUE_COLUMNS.get(tbl_name, set())

            # Table heading
            heading_text = f"{TABLE_SECTION_PREFIX}.{table_sequence}. {tbl_name}"
            h3 = doc.add_heading(heading_text, level=3)
            for run in h3.runs:
                run.font.name = FONT_NAME

            # Description
            desc = DESCRIPTIONS.get(tbl_name, "")
            if desc:
                p = doc.add_paragraph(desc)
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(11)
                    run.italic = True

            # Legend
            legend = doc.add_paragraph(
                "PK = Primary Key | FK = Foreign Key | UN = Unique | NN = Not Null"
            )
            legend.style = doc.styles["Normal"]
            for run in legend.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # Create table: No | Field | Type | PK | FK | UN | NN | Description
            num_rows = 1 + len(cols)
            tbl = doc.add_table(rows=num_rows, cols=8)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = True

            # Header row
            headers = ["No", "Field", "Type", "PK", "FK", "UN", "NN", "Description"]
            header_row = tbl.rows[0]
            for i, h_text in enumerate(headers):
                cell = header_row.cells[i]
                format_cell(cell, h_text, bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_shading(cell, HEADER_COLOR)
                set_cell_borders(cell)

            # Data rows
            for row_idx, col in enumerate(cols):
                row = tbl.rows[row_idx + 1]
                col_name = col["name"]
                is_pk = col["is_pk"]
                is_fk = col_name in fks
                # Case-sensitive unique check
                is_un = col_name in uniques or col_name.lower() in {u.lower() for u in uniques}
                is_nn = col["not_null"]

                # Get description
                field_desc = FIELD_DESCRIPTIONS.get((tbl_name, col_name), "")
                if not field_desc and col["comment"]:
                    field_desc = col["comment"]

                values = [
                    str(row_idx + 1),
                    col_name,
                    col["type"],
                    "V" if is_pk else "",
                    "V" if is_fk else "",
                    "V" if is_un else "",
                    "V" if is_nn else "",
                    field_desc,
                ]

                for ci, val in enumerate(values):
                    cell = row.cells[ci]
                    align = WD_ALIGN_PARAGRAPH.CENTER if ci in (0, 3, 4, 5, 6) else WD_ALIGN_PARAGRAPH.LEFT
                    format_cell(cell, val, bold=(ci == 1 and is_pk), size=10, alignment=align)
                    set_cell_borders(cell)

            # Add spacing after table
            doc.add_paragraph("")

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


# ===========================================================================
# Main
# ===========================================================================
def main():
    if not os.path.exists(SCHEMA_PATH):
        print(f"ERROR: Schema file not found: {SCHEMA_PATH}")
        sys.exit(1)

    print(f"Parsing SQL DDL from: {SCHEMA_PATH}")
    tables, pk_map, fk_map = parse_sql(SCHEMA_PATH)
    print(f"  Found {len(tables)} tables")

    total_cols = sum(len(c) for c in tables.values())
    total_fks = sum(len(f) for f in fk_map.values())
    print(f"  Total columns: {total_cols}")
    print(f"  Total FK relationships: {total_fks}")

    # Check for missing field descriptions
    missing = []
    for tbl_name, cols in tables.items():
        for col in cols:
            key = (tbl_name, col["name"])
            if key not in FIELD_DESCRIPTIONS and not col["comment"]:
                missing.append(key)
    if missing:
        print(f"\n  WARNING: {len(missing)} fields have no description:")
        for t, c in missing[:20]:
            print(f"    - {t}.{c}")
        if len(missing) > 20:
            print(f"    ... and {len(missing) - 20} more")

    print(f"\nGenerating Word document...")
    out = generate_docx(tables, pk_map, fk_map)
    print(f"  Output: {out}")
    print("Done!")


if __name__ == "__main__":
    main()
