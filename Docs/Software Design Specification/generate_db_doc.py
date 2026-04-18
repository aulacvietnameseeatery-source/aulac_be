#!/usr/bin/env python3
"""Generate database design documentation as Word (.docx) from PUML ERD definitions."""

import re
import os
import subprocess
import sys

# Auto-install python-docx if missing
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PUML_PATH = os.path.join(SCRIPT_DIR, "conceptual-erd.puml")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "database-design-tables.docx")

# ─────────────────────────────────────────────────────────
# Entity descriptions
# ─────────────────────────────────────────────────────────

DESCRIPTIONS = {
    "Customer": "Stores customer contact info (no login), loyalty status, and points.",
    "DishCategory": "Groups dishes into logical menu categories.",
    "Dish": "Represents menu items with price, category, and availability status.",
    "RestaurantTable": "Represents physical tables in the restaurant, including capacity, status, type, and QR image.",
    "Reservation": "Stores table reservations from online or phone sources.",
    "ReservationTable": "Junction table linking reservations to one or more restaurant tables.",
    "Order": "Represents dining orders created from tables, QR, web, or staff.",
    "OrderItem": "Stores individual line items within an order, referencing a dish with quantity and price.",
    "Payment": "Records payment transactions linked to orders.",
    "Tax": "Defines tax configurations that can be applied to orders.",
    "OrderCoupon": "Records coupons applied to a specific order with the calculated discount amount.",
    "OrderPromotion": "Records promotions applied to a specific order with the calculated discount amount.",
    "Supplier": "Stores supplier information for ingredients.",
    "Ingredient": "Master list of ingredients used for inventory tracking.",
    "CurrentStock": "Tracks the current on-hand quantity and minimum stock level for each ingredient.",
    "IngredientSupplier": "Junction table linking ingredients to their available suppliers.",
    "Recipe": "Defines the ingredient composition of a dish with quantity and unit.",
    "InventoryTransaction": "Records stock IN / OUT / ADJUST movements with supplier and purchase traceability.",
    "InventoryTransactionItem": "Stores individual line items within an inventory transaction.",
    "StaffAccount": "Stores all system users (Admin, Manager, Staff), authentication credentials, and account status.",
    "Role": "Defines system roles (Admin, Manager, Staff).",
    "Permission": "Defines screen + action permissions used for access control.",
    "RolePermission": "Junction table linking roles to their granted permissions.",
    "AuthSession": "Tracks active authentication sessions for staff users.",
    "LoginActivity": "Logs login/logout events for auditing and security monitoring.",
    "AuditLog": "Records staff actions on system resources for traceability.",
    "ServiceError": "Records staff service errors with severity, penalties, and resolution workflow.",
    "ServiceErrorCategory": "Defines types of service mistakes (wrong dish, delay, complaint, etc.).",
    "Promotion": "Defines promotions with validity period, type, usage limits, and activation status.",
    "PromotionRule": "Defines eligibility conditions for a promotion (min order value, required dish/category).",
    "PromotionTarget": "Defines which dishes or categories a promotion's discount applies to.",
    "Coupon": "Defines customer-owned discount coupons with usage limits and validity period.",
    "ShiftTemplate": "Defines reusable shift schedules with default start/end times and buffer minutes.",
    "ShiftAssignment": "Records planned work shifts assigned to staff members for a specific date.",
    "AttendanceRecord": "Tracks actual check-in/check-out times against planned shifts, including late and early-leave minutes.",
    "TimeLog": "Records individual punch-in/punch-out events with GPS and device info for an attendance record.",
    "Notification": "Stores system notifications with targeting rules, priority, and action metadata.",
    "NotificationReadState": "Tracks per-user read and acknowledgement status for each notification.",
    "NotificationPreference": "Stores per-user notification type preferences (enabled, sound).",
    "LookupType": "Defines lookup table categories (e.g. OrderStatus, TableType, Zone).",
    "LookupValue": "Stores individual values within a lookup type, used as configurable dropdown options system-wide.",
    "MediaAsset": "Stores metadata for images, 3D images, and videos using URL-based storage.",
    "DishMedium": "Junction table linking dishes to their media assets (images/videos).",
    "DishTag": "Associates keyword tags with dishes for filtering and promotion targeting.",
    "TableMedium": "Junction table linking restaurant tables to their media assets.",
    "InventoryTransactionMedium": "Junction table linking inventory transactions to their supporting media (receipts, photos).",
    "I18nLanguage": "Defines available languages for internationalization.",
    "I18nText": "Stores source text entries for translation.",
    "I18nTranslation": "Stores translated text for a specific language.",
    "SystemSetting": "Stores system-wide configuration key-value pairs for runtime behavior control.",
    "EmailTemplate": "Stores reusable email templates with subject and HTML body for system notifications.",
}

# ─────────────────────────────────────────────────────────
# Nullable properties per entity (from C# source analysis)
# ─────────────────────────────────────────────────────────

NULLABLE_MAP = {
    "AttendanceRecord": {"ActualCheckInAt", "ActualCheckOutAt", "AdjustmentReason", "ReviewedBy", "ReviewedAt"},
    "AuditLog": {"StaffId", "ActionCode", "TargetTable", "TargetId", "CreatedAt"},
    "AuthSession": {"CreatedAt", "Revoked", "DeviceInfo", "IpAddress"},
    "Coupon": {"CustomerId", "Description", "MaxUsage", "UsedCount", "CreatedAt"},
    "CurrentStock": {"LastUpdatedAt"},
    "Customer": {"FullName", "Email", "IsMember", "LoyaltyPoints", "CreatedAt"},
    "Dish": {"CreatedAt", "Description", "Slogan", "Note", "Calories", "ShortDescription", "DisplayOrder",
             "ChefRecommended", "PrepTimeMinutes", "CookTimeMinutes", "IsOnline"},
    "DishCategory": {"Description"},
    "DishMedium": {"IsPrimary"},
    "DishTag": set(),
    "EmailTemplate": {"Description", "UpdatedAt"},
    "I18nLanguage": {"IsActive"},
    "I18nText": {"Context", "CreatedAt", "UpdatedAt"},
    "I18nTranslation": {"UpdatedAt"},
    "Ingredient": {"TypeLvId", "CategoryLvId", "ImageId"},
    "IngredientSupplier": {"SupplierId", "IngredientId", "CreatedAt"},
    "InventoryTransaction": {"TransactionCode", "CreatedBy", "CreatedAt", "SubmittedAt", "ApprovedBy",
                              "ApprovedAt", "Note", "ExportReasonLvId", "StockCheckAreaNote", "SupplierId"},
    "InventoryTransactionItem": {"UnitPrice", "SystemQuantity", "ActualQuantity", "VarianceReasonLvId", "Note"},
    "InventoryTransactionMedium": {"IsPrimary"},
    "LoginActivity": {"SessionId", "DeviceInfo", "IpAddress"},
    "LookupType": {"Description", "IsSystem"},
    "LookupValue": {"Meta", "IsActive", "IsSystem", "Locked", "DeletedAt", "Description", "UpdateAt"},
    "MediaAsset": {"MimeType", "Width", "Height", "DurationSec", "CreatedAt"},
    "Notification": {"Body", "SoundKey", "ActionUrl", "EntityType", "EntityId", "MetadataJson",
                      "TargetPermissions", "TargetUserIds"},
    "NotificationPreference": set(),
    "NotificationReadState": {"ReadAt", "AcknowledgedAt"},
    "Order": {"TableId", "StaffId", "CreatedAt", "UpdatedAt", "TipAmount", "TaxId"},
    "OrderCoupon": {"AppliedAt"},
    "OrderItem": {"RejectReason", "Note"},
    "OrderPromotion": {"AppliedAt"},
    "Payment": {"PaidAt"},
    "Permission": set(),
    "Promotion": {"PromoCode", "Description", "MaxUsage", "UsedCount", "CreatedAt"},
    "PromotionRule": {"MinOrderValue", "MinQuantity", "RequiredDishId", "RequiredCategoryId"},
    "PromotionTarget": {"DishId", "CategoryId"},
    "Recipe": {"Note"},
    "Reservation": {"CustomerId", "Email", "CreatedAt", "Notes"},
    "RestaurantTable": {"IsOnline", "QrToken", "UpdatedByStaffId"},
    "Role": set(),
    "RolePermission": set(),
    "ServiceError": {"OrderId", "OrderItemId", "TableId", "PenaltyAmount", "IsResolved",
                      "ResolvedBy", "ResolvedAt", "CreatedAt"},
    "ServiceErrorCategory": {"Description"},
    "ShiftAssignment": {"Tags", "Notes"},
    "ShiftTemplate": {"Description", "BufferBeforeMinutes", "BufferAfterMinutes"},
    "StaffAccount": {"Phone", "Email", "LastLoginAt"},
    "Supplier": {"Phone", "Email", "Address", "TaxCode"},
    "SystemSetting": {"SettingName", "ValueString", "ValueInt", "ValueDecimal", "ValueBool",
                       "ValueJson", "Description", "UpdatedBy"},
    "TableMedium": {"IsPrimary"},
    "Tax": {"CreatedAt", "UpdatedAt"},
    "TimeLog": {"PunchOutTime", "GpsLocationIn", "GpsLocationOut", "DeviceIdIn", "DeviceIdOut"},
    "ReservationTable": set(),
}

# ─────────────────────────────────────────────────────────
# Unique fields per entity (inferred from business logic)
# ─────────────────────────────────────────────────────────

UNIQUE_MAP = {
    "Customer": {"phone"},
    "DishCategory": {"categoryName"},
    "RestaurantTable": {"tableCode", "qrToken"},
    "StaffAccount": {"username", "email"},
    "Role": {"roleCode"},
    "Permission": set(),  # composite unique on screenCode+actionCode
    "Supplier": {"taxCode"},
    "LookupType": {"typeCode"},
    "LookupValue": set(),  # composite unique typeId+valueCode
    "I18nLanguage": set(),  # langCode is PK
    "I18nText": {"textKey"},
    "SystemSetting": {"settingKey"},
    "EmailTemplate": {"templateCode"},
    "Promotion": {"promoCode"},
    "Coupon": {"couponCode"},
    "InventoryTransaction": {"transactionCode"},
    "ServiceErrorCategory": {"categoryCode"},
    "Notification": set(),
    "MediaAsset": set(),
}

# ─────────────────────────────────────────────────────────
# Field descriptions
# ─────────────────────────────────────────────────────────

FIELD_DESCRIPTIONS = {
    # Customer
    ("Customer", "customerId"): "Auto-incremented unique identifier for the customer",
    ("Customer", "fullName"): "Customer's full name",
    ("Customer", "phone"): "Customer's phone number, used as primary contact",
    ("Customer", "email"): "Customer's email address",
    ("Customer", "isMember"): "Whether the customer is a registered loyalty member",
    ("Customer", "loyaltyPoints"): "Accumulated loyalty points from orders",
    ("Customer", "createdAt"): "Timestamp when the customer record was created",

    # DishCategory
    ("DishCategory", "categoryId"): "Auto-incremented unique identifier",
    ("DishCategory", "categoryName"): "Display name of the dish category",
    ("DishCategory", "description"): "Optional description of the category",
    ("DishCategory", "displayOrder"): "Sort order for menu display",
    ("DishCategory", "isDisabled"): "Whether the category is hidden from the menu",

    # Dish
    ("Dish", "dishId"): "Auto-incremented unique identifier",
    ("Dish", "categoryId"): "References the dish category this dish belongs to",
    ("Dish", "dishName"): "Display name of the dish",
    ("Dish", "price"): "Current selling price",
    ("Dish", "description"): "Detailed description of the dish",
    ("Dish", "slogan"): "Marketing tagline for the dish",
    ("Dish", "note"): "Internal notes about the dish",
    ("Dish", "calories"): "Calorie count per serving",
    ("Dish", "shortDescription"): "Brief description for menus and cards",
    ("Dish", "displayOrder"): "Sort order within category",
    ("Dish", "chefRecommended"): "Whether the dish is flagged as chef's recommendation",
    ("Dish", "prepTimeMinutes"): "Estimated preparation time in minutes",
    ("Dish", "cookTimeMinutes"): "Estimated cooking time in minutes",
    ("Dish", "isOnline"): "Whether the dish is available for online ordering",
    ("Dish", "dishStatusLvId"): "References LookupValue for dish status (Available, Unavailable, etc.)",
    ("Dish", "createdAt"): "Timestamp when the dish was created",

    # RestaurantTable
    ("RestaurantTable", "tableId"): "Auto-incremented unique identifier",
    ("RestaurantTable", "tableCode"): "Unique short code displayed on the physical table (e.g. T01)",
    ("RestaurantTable", "capacity"): "Maximum seating capacity",
    ("RestaurantTable", "tableStatusLvId"): "References LookupValue for table status (Available, Occupied, Reserved)",
    ("RestaurantTable", "tableTypeLvId"): "References LookupValue for table type (Indoor, Outdoor, VIP)",
    ("RestaurantTable", "zoneLvId"): "References LookupValue for restaurant zone",
    ("RestaurantTable", "isOnline"): "Whether the table is available for online reservation",
    ("RestaurantTable", "qrToken"): "Unique token embedded in QR code for self-service ordering",
    ("RestaurantTable", "isDeleted"): "Soft-delete flag",
    ("RestaurantTable", "createdAt"): "Timestamp when the table was created",
    ("RestaurantTable", "updatedAt"): "Timestamp of the last update",

    # Reservation
    ("Reservation", "reservationId"): "Auto-incremented unique identifier",
    ("Reservation", "customerId"): "References the customer who made the reservation",
    ("Reservation", "customerName"): "Name provided for the reservation (may differ from customer record)",
    ("Reservation", "phone"): "Contact phone for the reservation",
    ("Reservation", "email"): "Contact email for the reservation",
    ("Reservation", "partySize"): "Number of guests expected",
    ("Reservation", "reservedTime"): "Date and time of the reservation",
    ("Reservation", "reservationStatusLvId"): "References LookupValue for reservation status (Pending, Confirmed, Cancelled)",
    ("Reservation", "sourceLvId"): "References LookupValue for reservation source (Online, Phone, Walk-in)",
    ("Reservation", "notes"): "Additional notes or special requests",
    ("Reservation", "createdAt"): "Timestamp when the reservation was created",

    # ReservationTable
    ("ReservationTable", "reservationId"): "References the reservation",
    ("ReservationTable", "tableId"): "References the assigned restaurant table",

    # Order
    ("Order", "orderId"): "Auto-incremented unique identifier",
    ("Order", "tableId"): "References the table where the order is served",
    ("Order", "staffId"): "References the staff member handling the order",
    ("Order", "customerId"): "References the customer placing the order",
    ("Order", "orderStatusLvId"): "References LookupValue for order status (New, Preparing, Completed, Cancelled)",
    ("Order", "sourceLvId"): "References LookupValue for order source (POS, QR, Web)",
    ("Order", "taxId"): "References the applicable tax configuration",
    ("Order", "totalAmount"): "Final total after tax and discounts",
    ("Order", "subTotalAmount"): "Subtotal before tax and discounts",
    ("Order", "tipAmount"): "Optional tip amount",
    ("Order", "taxAmount"): "Calculated tax amount",
    ("Order", "createdAt"): "Timestamp when the order was created",
    ("Order", "updatedAt"): "Timestamp of the last update",

    # OrderItem
    ("OrderItem", "orderItemId"): "Auto-incremented unique identifier",
    ("OrderItem", "orderId"): "References the parent order",
    ("OrderItem", "dishId"): "References the ordered dish",
    ("OrderItem", "quantity"): "Number of servings ordered",
    ("OrderItem", "price"): "Unit price at the time of ordering",
    ("OrderItem", "itemStatusLvId"): "References LookupValue for item status (Pending, Preparing, Served, Rejected)",
    ("OrderItem", "rejectReason"): "Reason for rejection if item was rejected",
    ("OrderItem", "note"): "Special instructions for the kitchen",

    # Payment
    ("Payment", "paymentId"): "Auto-incremented unique identifier",
    ("Payment", "orderId"): "References the order being paid for",
    ("Payment", "methodLvId"): "References LookupValue for payment method (Cash, Card, Transfer)",
    ("Payment", "receivedAmount"): "Amount received from the customer",
    ("Payment", "changeAmount"): "Change returned to the customer",
    ("Payment", "paidAt"): "Timestamp when the payment was made",

    # Tax
    ("Tax", "taxId"): "Auto-incremented unique identifier",
    ("Tax", "taxName"): "Display name of the tax (e.g. VAT 10%)",
    ("Tax", "taxRate"): "Tax rate as decimal (e.g. 0.10 for 10%)",
    ("Tax", "taxType"): "Type of tax calculation (Percentage, Fixed)",
    ("Tax", "isActive"): "Whether the tax is currently active",
    ("Tax", "isDefault"): "Whether this is the default tax applied to new orders",
    ("Tax", "createdAt"): "Timestamp when the tax was created",

    # OrderCoupon
    ("OrderCoupon", "orderCouponId"): "Auto-incremented unique identifier",
    ("OrderCoupon", "orderId"): "References the order",
    ("OrderCoupon", "couponId"): "References the applied coupon",
    ("OrderCoupon", "discountAmount"): "Calculated discount amount from the coupon",
    ("OrderCoupon", "appliedAt"): "Timestamp when the coupon was applied",

    # OrderPromotion
    ("OrderPromotion", "orderPromotionId"): "Auto-incremented unique identifier",
    ("OrderPromotion", "orderId"): "References the order",
    ("OrderPromotion", "promotionId"): "References the applied promotion",
    ("OrderPromotion", "discountAmount"): "Calculated discount amount from the promotion",
    ("OrderPromotion", "appliedAt"): "Timestamp when the promotion was applied",

    # Supplier
    ("Supplier", "supplierId"): "Auto-incremented unique identifier",
    ("Supplier", "supplierName"): "Business name of the supplier",
    ("Supplier", "phone"): "Contact phone number",
    ("Supplier", "email"): "Contact email address",
    ("Supplier", "address"): "Business address",
    ("Supplier", "taxCode"): "Tax identification code",

    # Ingredient
    ("Ingredient", "ingredientId"): "Auto-incremented unique identifier",
    ("Ingredient", "ingredientName"): "Name of the ingredient",
    ("Ingredient", "unitLvId"): "References LookupValue for measurement unit (kg, g, l, piece)",
    ("Ingredient", "typeLvId"): "References LookupValue for ingredient type",
    ("Ingredient", "categoryLvId"): "References LookupValue for ingredient category (Dairy, Meat, Produce)",
    ("Ingredient", "imageId"): "References MediaAsset for the ingredient image",

    # CurrentStock
    ("CurrentStock", "ingredientId"): "References the ingredient (1-to-1 relationship)",
    ("CurrentStock", "quantityOnHand"): "Current available quantity in stock",
    ("CurrentStock", "minStockLevel"): "Minimum quantity threshold for low-stock alerts",
    ("CurrentStock", "lastUpdatedAt"): "Timestamp of the last stock update",

    # IngredientSupplier
    ("IngredientSupplier", "ingredientSupplierId"): "Auto-incremented unique identifier",
    ("IngredientSupplier", "supplierId"): "References the supplier",
    ("IngredientSupplier", "ingredientId"): "References the ingredient",
    ("IngredientSupplier", "createdAt"): "Timestamp when the link was created",

    # Recipe
    ("Recipe", "dishId"): "References the dish this recipe belongs to",
    ("Recipe", "ingredientId"): "References the ingredient used",
    ("Recipe", "quantity"): "Amount of ingredient required per serving",
    ("Recipe", "unit"): "Unit of measurement for the quantity",
    ("Recipe", "note"): "Preparation notes for this ingredient",

    # InventoryTransaction
    ("InventoryTransaction", "transactionId"): "Auto-incremented unique identifier",
    ("InventoryTransaction", "transactionCode"): "Unique human-readable transaction code",
    ("InventoryTransaction", "createdBy"): "References the staff who created the transaction",
    ("InventoryTransaction", "approvedBy"): "References the staff who approved the transaction",
    ("InventoryTransaction", "typeLvId"): "References LookupValue for transaction type (Import, Export, StockCheck)",
    ("InventoryTransaction", "statusLvId"): "References LookupValue for transaction status (Draft, Submitted, Approved)",
    ("InventoryTransaction", "supplierId"): "References the supplier for import transactions",
    ("InventoryTransaction", "exportReasonLvId"): "References LookupValue for export reason",
    ("InventoryTransaction", "stockCheckAreaNote"): "Area description for stock-check transactions",
    ("InventoryTransaction", "note"): "Additional notes about the transaction",
    ("InventoryTransaction", "createdAt"): "Timestamp when the transaction was created",
    ("InventoryTransaction", "submittedAt"): "Timestamp when submitted for approval",
    ("InventoryTransaction", "approvedAt"): "Timestamp when approved",

    # InventoryTransactionItem
    ("InventoryTransactionItem", "transactionItemId"): "Auto-incremented unique identifier",
    ("InventoryTransactionItem", "transactionId"): "References the parent transaction",
    ("InventoryTransactionItem", "ingredientId"): "References the ingredient",
    ("InventoryTransactionItem", "quantity"): "Quantity moved in this line item",
    ("InventoryTransactionItem", "unitLvId"): "References LookupValue for measurement unit",
    ("InventoryTransactionItem", "unitPrice"): "Unit price for import transactions",
    ("InventoryTransactionItem", "systemQuantity"): "System-recorded quantity (for stock-check comparison)",
    ("InventoryTransactionItem", "actualQuantity"): "Actual counted quantity (for stock-check)",
    ("InventoryTransactionItem", "varianceReasonLvId"): "References LookupValue for variance reason",
    ("InventoryTransactionItem", "note"): "Additional notes about this line item",

    # StaffAccount
    ("StaffAccount", "accountId"): "Auto-incremented unique identifier",
    ("StaffAccount", "fullName"): "Staff member's full name",
    ("StaffAccount", "phone"): "Contact phone number",
    ("StaffAccount", "email"): "Contact email address",
    ("StaffAccount", "username"): "Unique login username",
    ("StaffAccount", "passwordHash"): "Hashed password for authentication",
    ("StaffAccount", "roleId"): "References the assigned role",
    ("StaffAccount", "accountStatusLvId"): "References LookupValue for account status (Active, Inactive, Locked)",
    ("StaffAccount", "isLocked"): "Whether the account is temporarily locked",
    ("StaffAccount", "lastLoginAt"): "Timestamp of the last successful login",
    ("StaffAccount", "createdAt"): "Timestamp when the account was created",

    # Role
    ("Role", "roleId"): "Auto-incremented unique identifier",
    ("Role", "roleCode"): "Unique code identifying the role (e.g. ADMIN, MANAGER, STAFF)",
    ("Role", "roleName"): "Display name of the role",
    ("Role", "roleStatusLvId"): "References LookupValue for role status (Active, Inactive)",

    # Permission
    ("Permission", "permissionId"): "Auto-incremented unique identifier",
    ("Permission", "screenCode"): "Code identifying the UI screen or module",
    ("Permission", "actionCode"): "Code identifying the allowed action (View, Create, Edit, Delete)",

    # RolePermission
    ("RolePermission", "roleId"): "References the role",
    ("RolePermission", "permissionId"): "References the permission",

    # AuthSession
    ("AuthSession", "sessionId"): "Auto-incremented unique identifier",
    ("AuthSession", "userId"): "References the staff account",
    ("AuthSession", "tokenHash"): "Hashed refresh token for session validation",
    ("AuthSession", "expiresAt"): "Token expiry timestamp",
    ("AuthSession", "revoked"): "Whether the session has been revoked",
    ("AuthSession", "deviceInfo"): "User-Agent or device description",
    ("AuthSession", "ipAddress"): "Client IP address at session creation",
    ("AuthSession", "createdAt"): "Timestamp when the session was created",

    # LoginActivity
    ("LoginActivity", "loginActivityId"): "Auto-incremented unique identifier",
    ("LoginActivity", "staffId"): "References the staff account",
    ("LoginActivity", "sessionId"): "References the associated auth session",
    ("LoginActivity", "eventType"): "Type of event (Login, Logout, TokenRefresh, FailedLogin)",
    ("LoginActivity", "deviceInfo"): "User-Agent or device description",
    ("LoginActivity", "ipAddress"): "Client IP address",
    ("LoginActivity", "occurredAt"): "Timestamp when the event occurred",

    # AuditLog
    ("AuditLog", "logId"): "Auto-incremented unique identifier",
    ("AuditLog", "staffId"): "References the staff who performed the action",
    ("AuditLog", "actionCode"): "Code describing the action (e.g. CREATE, UPDATE, DELETE)",
    ("AuditLog", "targetTable"): "Name of the affected database table",
    ("AuditLog", "targetId"): "Primary key of the affected record",
    ("AuditLog", "createdAt"): "Timestamp when the action occurred",

    # ServiceError
    ("ServiceError", "errorId"): "Auto-incremented unique identifier",
    ("ServiceError", "staffId"): "References the staff member responsible for the error",
    ("ServiceError", "orderId"): "References the related order (if applicable)",
    ("ServiceError", "orderItemId"): "References the related order item (if applicable)",
    ("ServiceError", "tableId"): "References the related table (if applicable)",
    ("ServiceError", "categoryId"): "References the service error category",
    ("ServiceError", "severityLvId"): "References LookupValue for severity level (Low, Medium, High, Critical)",
    ("ServiceError", "description"): "Detailed description of the service error",
    ("ServiceError", "penaltyAmount"): "Monetary penalty amount (if applicable)",
    ("ServiceError", "isResolved"): "Whether the error has been resolved",
    ("ServiceError", "resolvedBy"): "References the staff member who resolved the error",
    ("ServiceError", "resolvedAt"): "Timestamp when the error was resolved",
    ("ServiceError", "createdAt"): "Timestamp when the error was recorded",

    # ServiceErrorCategory
    ("ServiceErrorCategory", "categoryId"): "Auto-incremented unique identifier",
    ("ServiceErrorCategory", "categoryCode"): "Unique code for the category",
    ("ServiceErrorCategory", "categoryName"): "Display name of the category",
    ("ServiceErrorCategory", "description"): "Optional description of the error category",

    # Promotion
    ("Promotion", "promotionId"): "Auto-incremented unique identifier",
    ("Promotion", "promoCode"): "Unique promotion code",
    ("Promotion", "promoName"): "Display name of the promotion",
    ("Promotion", "description"): "Detailed description of the promotion",
    ("Promotion", "typeLvId"): "References LookupValue for promotion type (Percentage, Fixed, BuyXGetY)",
    ("Promotion", "promotionStatusLvId"): "References LookupValue for promotion status (Active, Inactive, Expired)",
    ("Promotion", "discountValue"): "Discount value (percentage or fixed amount depending on type)",
    ("Promotion", "maxUsage"): "Maximum number of times the promotion can be used",
    ("Promotion", "usedCount"): "Current number of times the promotion has been used",
    ("Promotion", "startTime"): "Promotion validity start date/time",
    ("Promotion", "endTime"): "Promotion validity end date/time",
    ("Promotion", "createdAt"): "Timestamp when the promotion was created",

    # PromotionRule
    ("PromotionRule", "ruleId"): "Auto-incremented unique identifier",
    ("PromotionRule", "promotionId"): "References the parent promotion",
    ("PromotionRule", "minOrderValue"): "Minimum order subtotal required to qualify",
    ("PromotionRule", "minQuantity"): "Minimum item quantity required to qualify",
    ("PromotionRule", "requiredDishId"): "Specific dish required in the order to qualify",
    ("PromotionRule", "requiredCategoryId"): "Specific dish category required in the order to qualify",

    # PromotionTarget
    ("PromotionTarget", "targetId"): "Auto-incremented unique identifier",
    ("PromotionTarget", "promotionId"): "References the parent promotion",
    ("PromotionTarget", "dishId"): "Specific dish the discount applies to",
    ("PromotionTarget", "categoryId"): "Specific category the discount applies to",

    # Coupon
    ("Coupon", "couponId"): "Auto-incremented unique identifier",
    ("Coupon", "customerId"): "References the customer who owns the coupon",
    ("Coupon", "couponCode"): "Unique coupon code",
    ("Coupon", "couponName"): "Display name of the coupon",
    ("Coupon", "description"): "Detailed description",
    ("Coupon", "typeLvId"): "References LookupValue for coupon type (Percentage, Fixed)",
    ("Coupon", "couponStatusLvId"): "References LookupValue for coupon status (Active, Used, Expired)",
    ("Coupon", "discountValue"): "Discount value (percentage or fixed amount)",
    ("Coupon", "maxUsage"): "Maximum number of times the coupon can be used",
    ("Coupon", "usedCount"): "Current number of times the coupon has been used",
    ("Coupon", "startTime"): "Coupon validity start date/time",
    ("Coupon", "endTime"): "Coupon validity end date/time",
    ("Coupon", "createdAt"): "Timestamp when the coupon was created",

    # ShiftTemplate
    ("ShiftTemplate", "shiftTemplateId"): "Auto-incremented unique identifier",
    ("ShiftTemplate", "templateName"): "Display name of the shift template (e.g. Morning, Evening)",
    ("ShiftTemplate", "defaultStartTime"): "Default shift start time",
    ("ShiftTemplate", "defaultEndTime"): "Default shift end time",
    ("ShiftTemplate", "description"): "Optional description of the shift",
    ("ShiftTemplate", "bufferBeforeMinutes"): "Allowed early check-in buffer in minutes",
    ("ShiftTemplate", "bufferAfterMinutes"): "Allowed late check-out buffer in minutes",
    ("ShiftTemplate", "isActive"): "Whether the template is currently active",
    ("ShiftTemplate", "createdBy"): "References the staff who created the template",
    ("ShiftTemplate", "createdAt"): "Timestamp when the template was created",

    # ShiftAssignment
    ("ShiftAssignment", "shiftAssignmentId"): "Auto-incremented unique identifier",
    ("ShiftAssignment", "shiftTemplateId"): "References the shift template used",
    ("ShiftAssignment", "staffId"): "References the assigned staff member",
    ("ShiftAssignment", "workDate"): "The scheduled work date",
    ("ShiftAssignment", "plannedStartAt"): "Planned shift start date/time",
    ("ShiftAssignment", "plannedEndAt"): "Planned shift end date/time",
    ("ShiftAssignment", "assignmentStatusLvId"): "References LookupValue for assignment status (Scheduled, Completed, Absent)",
    ("ShiftAssignment", "isActive"): "Whether the assignment is active",
    ("ShiftAssignment", "tags"): "Optional tags for categorization",
    ("ShiftAssignment", "notes"): "Optional notes about the assignment",
    ("ShiftAssignment", "assignedBy"): "References the staff who created the assignment",
    ("ShiftAssignment", "assignedAt"): "Timestamp when the assignment was created",

    # AttendanceRecord
    ("AttendanceRecord", "attendanceId"): "Auto-incremented unique identifier",
    ("AttendanceRecord", "shiftAssignmentId"): "References the associated shift assignment",
    ("AttendanceRecord", "attendanceStatusLvId"): "References LookupValue for attendance status (Present, Late, Absent)",
    ("AttendanceRecord", "actualCheckInAt"): "Actual check-in timestamp",
    ("AttendanceRecord", "actualCheckOutAt"): "Actual check-out timestamp",
    ("AttendanceRecord", "lateMinutes"): "Number of minutes late for check-in",
    ("AttendanceRecord", "earlyLeaveMinutes"): "Number of minutes early for check-out",
    ("AttendanceRecord", "workedMinutes"): "Total minutes worked in the shift",
    ("AttendanceRecord", "isManualAdjustment"): "Whether the record was manually adjusted by a manager",
    ("AttendanceRecord", "adjustmentReason"): "Reason for manual adjustment (if applicable)",
    ("AttendanceRecord", "reviewedBy"): "References the manager who reviewed the record",
    ("AttendanceRecord", "reviewedAt"): "Timestamp when the record was reviewed",

    # TimeLog
    ("TimeLog", "timeLogId"): "Auto-incremented unique identifier",
    ("TimeLog", "attendanceRecordId"): "References the associated attendance record",
    ("TimeLog", "punchInTime"): "Punch-in timestamp",
    ("TimeLog", "punchOutTime"): "Punch-out timestamp",
    ("TimeLog", "gpsLocationIn"): "GPS coordinates at punch-in",
    ("TimeLog", "gpsLocationOut"): "GPS coordinates at punch-out",
    ("TimeLog", "deviceIdIn"): "Device identifier at punch-in",
    ("TimeLog", "deviceIdOut"): "Device identifier at punch-out",
    ("TimeLog", "validationStatus"): "Validation result (Valid, Suspicious, Invalid)",
    ("TimeLog", "punchDurationMinutes"): "Duration between punch-in and punch-out in minutes",

    # Notification
    ("Notification", "notificationId"): "Auto-incremented unique identifier",
    ("Notification", "type"): "Notification type code (e.g. ORDER_NEW, SHIFT_REMINDER)",
    ("Notification", "title"): "Notification title displayed to users",
    ("Notification", "body"): "Notification body text",
    ("Notification", "priority"): "Priority level (Normal, High, Urgent)",
    ("Notification", "requireAck"): "Whether the notification requires user acknowledgement",
    ("Notification", "soundKey"): "Sound identifier for the notification",
    ("Notification", "actionUrl"): "Deep-link URL for the notification action",
    ("Notification", "entityType"): "Type of the related entity (e.g. Order, Reservation)",
    ("Notification", "entityId"): "ID of the related entity",
    ("Notification", "metadataJson"): "Additional metadata as JSON string",
    ("Notification", "targetPermissions"): "Comma-separated permission codes for targeting",
    ("Notification", "targetUserIds"): "Comma-separated specific user IDs for targeting",
    ("Notification", "createdAt"): "Timestamp when the notification was created",

    # NotificationReadState
    ("NotificationReadState", "notificationReadStateId"): "Auto-incremented unique identifier",
    ("NotificationReadState", "notificationId"): "References the notification",
    ("NotificationReadState", "userId"): "ID of the user",
    ("NotificationReadState", "isRead"): "Whether the user has read the notification",
    ("NotificationReadState", "isAcknowledged"): "Whether the user has acknowledged the notification",
    ("NotificationReadState", "readAt"): "Timestamp when marked as read",
    ("NotificationReadState", "acknowledgedAt"): "Timestamp when acknowledged",

    # NotificationPreference
    ("NotificationPreference", "notificationPreferenceId"): "Auto-incremented unique identifier",
    ("NotificationPreference", "userId"): "ID of the user",
    ("NotificationPreference", "notificationType"): "Notification type code",
    ("NotificationPreference", "isEnabled"): "Whether notifications of this type are enabled",
    ("NotificationPreference", "soundEnabled"): "Whether sound is enabled for this type",

    # LookupType
    ("LookupType", "typeId"): "Auto-incremented unique identifier",
    ("LookupType", "typeCode"): "Unique code for the lookup type (e.g. OrderStatus, TableType)",
    ("LookupType", "typeName"): "Display name of the lookup type",
    ("LookupType", "description"): "Optional description",
    ("LookupType", "isConfigurable"): "Whether values can be modified by users",
    ("LookupType", "isSystem"): "Whether this is a system-managed lookup type",

    # LookupValue
    ("LookupValue", "valueId"): "Auto-incremented unique identifier",
    ("LookupValue", "typeId"): "References the parent lookup type",
    ("LookupValue", "valueCode"): "Unique code within the type",
    ("LookupValue", "valueName"): "Display name of the value",
    ("LookupValue", "sortOrder"): "Sort order for display",
    ("LookupValue", "isActive"): "Whether the value is currently active",
    ("LookupValue", "meta"): "Optional metadata as JSON string",
    ("LookupValue", "isSystem"): "Whether this is a system-managed value",
    ("LookupValue", "locked"): "Whether the value is locked from editing",

    # MediaAsset
    ("MediaAsset", "mediaId"): "Auto-incremented unique identifier",
    ("MediaAsset", "url"): "URL path to the stored media file",
    ("MediaAsset", "mimeType"): "MIME type of the media (e.g. image/jpeg, video/mp4)",
    ("MediaAsset", "width"): "Image/video width in pixels",
    ("MediaAsset", "height"): "Image/video height in pixels",
    ("MediaAsset", "durationSec"): "Video duration in seconds",
    ("MediaAsset", "mediaTypeLvId"): "References LookupValue for media type (Image, 3DImage, Video)",
    ("MediaAsset", "createdAt"): "Timestamp when the media was uploaded",

    # DishMedium
    ("DishMedium", "dishId"): "References the dish",
    ("DishMedium", "mediaId"): "References the media asset",
    ("DishMedium", "isPrimary"): "Whether this is the primary/thumbnail image for the dish",

    # DishTag
    ("DishTag", "dishTagId"): "Auto-incremented unique identifier",
    ("DishTag", "dishId"): "References the dish",
    ("DishTag", "tagId"): "References LookupValue for the tag",

    # TableMedium
    ("TableMedium", "tableId"): "References the restaurant table",
    ("TableMedium", "mediaId"): "References the media asset",
    ("TableMedium", "isPrimary"): "Whether this is the primary image for the table",

    # InventoryTransactionMedium
    ("InventoryTransactionMedium", "transactionId"): "References the inventory transaction",
    ("InventoryTransactionMedium", "mediaId"): "References the media asset",
    ("InventoryTransactionMedium", "isPrimary"): "Whether this is the primary receipt/photo",

    # I18nLanguage
    ("I18nLanguage", "langCode"): "ISO language code (e.g. en, vi, zh)",
    ("I18nLanguage", "langName"): "Display name of the language",
    ("I18nLanguage", "isActive"): "Whether the language is currently enabled",

    # I18nText
    ("I18nText", "textId"): "Auto-incremented unique identifier",
    ("I18nText", "textKey"): "Unique key identifying the text entry",
    ("I18nText", "sourceLangCode"): "Language code of the original source text",
    ("I18nText", "sourceText"): "Original text in the source language",
    ("I18nText", "context"): "Context hint for translators",

    # I18nTranslation
    ("I18nTranslation", "textId"): "References the source text entry",
    ("I18nTranslation", "langCode"): "References the target language",
    ("I18nTranslation", "translatedText"): "Translated text",
    ("I18nTranslation", "updatedAt"): "Timestamp of the last translation update",

    # SystemSetting
    ("SystemSetting", "settingId"): "Auto-incremented unique identifier",
    ("SystemSetting", "settingKey"): "Unique key identifying the setting",
    ("SystemSetting", "settingName"): "Display name of the setting",
    ("SystemSetting", "valueType"): "Data type of the value (String, Int, Decimal, Bool, Json)",
    ("SystemSetting", "valueString"): "String value (if valueType is String)",
    ("SystemSetting", "valueInt"): "Integer value (if valueType is Int)",
    ("SystemSetting", "valueDecimal"): "Decimal value (if valueType is Decimal)",
    ("SystemSetting", "valueBool"): "Boolean value (if valueType is Bool)",
    ("SystemSetting", "valueJson"): "JSON value (if valueType is Json)",
    ("SystemSetting", "description"): "Description of what this setting controls",
    ("SystemSetting", "isSensitive"): "Whether the value should be masked in UI",
    ("SystemSetting", "updatedBy"): "References the staff who last updated the setting",

    # EmailTemplate
    ("EmailTemplate", "templateId"): "Auto-incremented unique identifier",
    ("EmailTemplate", "templateCode"): "Unique code identifying the template",
    ("EmailTemplate", "templateName"): "Display name of the template",
    ("EmailTemplate", "subject"): "Email subject line (may contain placeholders)",
    ("EmailTemplate", "bodyHtml"): "Email body in HTML format (may contain placeholders)",
    ("EmailTemplate", "description"): "Description of when this template is used",
}


def parse_puml(filepath):
    """Parse PlantUML file into ordered list of entities."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    entities = []
    seen = set()

    for match in re.finditer(
        r'entity\s+"([^"]+)"\s+as\s+(\w+)\s+<<(\w+)>>\s*\{([^}]*)\}',
        content, re.DOTALL
    ):
        name, alias, domain, body = match.group(1), match.group(2), match.group(3), match.group(4)
        if name in seen:
            continue
        seen.add(name)

        fields = []
        for line in body.strip().split('\n'):
            line = line.strip()
            if not line or line == '--':
                continue
            fm = re.match(
                r'(\*)?\s*(\w+)\s*:\s*(\w+)(?:\s+<<([^>]+)>>)?(?:\s+<<([^>]+)>>)?',
                line
            )
            if fm:
                is_pk_star = fm.group(1) == '*'
                fname = fm.group(2)
                ftype = fm.group(3)
                markers = set()
                for g in [fm.group(4), fm.group(5)]:
                    if g:
                        for part in g.split(','):
                            part = part.strip()
                            if part in ('PK', 'FK'):
                                markers.add(part)
                if is_pk_star and 'PK' not in markers:
                    markers.add('PK')
                fields.append({
                    'name': fname,
                    'type': ftype,
                    'is_pk': 'PK' in markers,
                    'is_fk': 'FK' in markers,
                })

        entities.append({'name': name, 'alias': alias, 'domain': domain, 'fields': fields})

    return entities


def is_nullable(entity_name, field_name):
    """Check if a field is nullable based on C# source analysis."""
    nullables = NULLABLE_MAP.get(entity_name, set())
    # Case-insensitive match
    for n in nullables:
        if n.lower() == field_name.lower():
            return True
    return False


def is_unique(entity_name, field_name):
    """Check if a field has a unique constraint."""
    uniques = UNIQUE_MAP.get(entity_name, set())
    for u in uniques:
        if u.lower() == field_name.lower():
            return True
    return False


def get_description(entity_name, field_name):
    """Get field description."""
    desc = FIELD_DESCRIPTIONS.get((entity_name, field_name))
    if desc:
        return desc
    return ""


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = tcBorders.makeelement(qn(f'w:{edge}'), {
                qn('w:val'): val.get('val', 'single'),
                qn('w:sz'): str(val.get('sz', 4)),
                qn('w:space'): '0',
                qn('w:color'): val.get('color', '000000'),
            })
            tcBorders.append(el)
    tcPr.append(tcBorders)


def format_cell(cell, text, bold=False, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Format a table cell with consistent styling."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.bold = bold
    # Reduce cell margins
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def generate_docx(entities):
    """Generate Word document with database design tables."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Group by domain
    domain_order = ['core', 'order', 'inventory', 'staff', 'promo', 'shift', 'notification', 'infra']
    domain_labels = {
        'core': 'Core Domain',
        'order': 'Order & Payment',
        'inventory': 'Inventory & Supply Chain',
        'staff': 'Staff, Authentication & Audit',
        'promo': 'Promotion & Coupon',
        'shift': 'Shift & Attendance',
        'notification': 'Notification',
        'infra': 'Infrastructure',
    }

    by_domain = {}
    for ent in entities:
        d = ent['domain']
        if d not in by_domain:
            by_domain[d] = []
        by_domain[d].append(ent)

    for di, domain in enumerate(domain_order):
        if domain not in by_domain:
            continue

        for ei, ent in enumerate(by_domain[domain]):
            ename = ent['name']
            desc = DESCRIPTIONS.get(ename, "")
            subsection = f"1.3.{di + 1}.{ei + 1}"

            # ── Section heading: "1.3.x.y table_name" ──
            heading = doc.add_heading(level=3)
            run = heading.add_run(f"{subsection} {ename}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = True
            run.italic = True

            # ── Description paragraph ──
            p = doc.add_paragraph()
            run = p.add_run(desc)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

            # ── Legend line ──
            p2 = doc.add_paragraph()
            run2 = p2.add_run("* PK~Primary Key; FK~Foreign Key; UN~Unique; NN ~ not null")
            run2.font.name = "Times New Roman"
            run2.font.size = Pt(10)
            run2.italic = True

            # ── Field table ──
            fields = ent['fields']
            num_rows = len(fields) + 1  # +1 for header
            table = doc.add_table(rows=num_rows, cols=7)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'

            # Column widths (approximate)
            col_widths = [Cm(1.0), Cm(4.5), Cm(1.0), Cm(1.0), Cm(1.0), Cm(1.0), Cm(6.5)]

            # Header row
            headers = ["No", "Field", "PK", "FK", "UN", "NN", "Description"]
            header_row = table.rows[0]
            for ci, h in enumerate(headers):
                cell = header_row.cells[ci]
                format_cell(cell, h, bold=True, font_size=10,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_shading(cell, "FBE4D5")  # light peach like template
                cell.width = col_widths[ci]

            # Data rows
            for fi, field in enumerate(fields):
                row = table.rows[fi + 1]
                fname = field['name']
                pk = "✓" if field['is_pk'] else ""
                fk = "✓" if field['is_fk'] else ""
                un = "✓" if is_unique(ename, fname) else ""

                nullable = is_nullable(ename, fname)
                nn = ""
                if field['is_pk']:
                    nn = "✓"
                elif not nullable:
                    nn = "✓"

                fdesc = get_description(ename, fname)

                values = [f"{fi + 1:02d}", fname, pk, fk, un, nn, fdesc]
                for ci, val in enumerate(values):
                    cell = row.cells[ci]
                    align = WD_ALIGN_PARAGRAPH.CENTER if ci in (0, 2, 3, 4, 5) else WD_ALIGN_PARAGRAPH.LEFT
                    format_cell(cell, val, font_size=10, align=align)
                    cell.width = col_widths[ci]

            # Add spacing after table
            doc.add_paragraph("")

    return doc


def main():
    print(f"Reading: {PUML_PATH}")
    entities = parse_puml(PUML_PATH)
    print(f"Parsed {len(entities)} entities")

    doc = generate_docx(entities)
    doc.save(OUTPUT_PATH)
    print(f"Wrote: {OUTPUT_PATH} ({os.path.getsize(OUTPUT_PATH):,} bytes)")

    # Stats
    total_fields = sum(len(e['fields']) for e in entities)
    descs = sum(1 for e in entities for f in e['fields']
                if get_description(e['name'], f['name']))
    print(f"Total fields: {total_fields}, with descriptions: {descs}")


if __name__ == '__main__':
    main()
